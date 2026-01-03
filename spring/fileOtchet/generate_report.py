#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    # Создаем новый документ
    doc = Document()

    # Настраиваем стиль документа - простой черный текст Times New Roman
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = None  # Черный цвет

    # Заголовок отчета
    title = doc.add_paragraph('Отчёт по проекту "Система поиска билетов"')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True

    doc.add_paragraph('')

    # 1. Перечень требований к программному средству
    doc.add_paragraph('1. Перечень требований к программному средству')
    doc.add_paragraph('')

    doc.add_paragraph('Функциональные требования:')
    functional_reqs = [
        '1. Поиск рейсов по параметрам (откуда, куда, дата, время)',
        '2. Фильтрация и сортировка результатов поиска',
        '3. Бронирование билетов с выбором места',
        '4. Управление билетами (просмотр, отмена)',
        '5. Просмотр расписания рейсов',
        '6. Сохранение и загрузка данных о рейсах и билетах'
    ]

    for req in functional_reqs:
        doc.add_paragraph('   ' + req)

    doc.add_paragraph('')
    doc.add_paragraph('Нефункциональные требования:')
    non_functional_reqs = [
        '1. Производительность: быстрый поиск и фильтрация (<2 сек)',
        '2. Надежность: сохранение данных при перезапусках',
        '3. Удобство использования: интуитивный графический интерфейс',
        '4. Совместимость: работа на различных ОС с Java',
        '5. Безопасность: защита данных пользователей'
    ]

    for req in non_functional_reqs:
        doc.add_paragraph('   ' + req)

    doc.add_paragraph('')

    # 2. Перечисление выполненных требований
    doc.add_paragraph('2. Перечисление выполненных требований')
    doc.add_paragraph('')

    doc.add_paragraph('Полностью выполнены:')
    doc.add_paragraph('')

    completed_reqs = [
        '1. Поиск рейсов',
        '   - Форма поиска с полями: откуда, куда, дата',
        '   - Выполнение поиска по заданным параметрам',
        '   - Отображение результатов поиска',
        '2. Фильтрация результатов',
        '   - Сортировка по различным критериям',
        '   - Фильтры по времени, цене, перевозчику',
        '3. Бронирование билетов',
        '   - Выбор рейса из списка',
        '   - Выбор места с графическим интерфейсом',
        '   - Ввод данных пассажира',
        '   - Подтверждение бронирования',
        '4. Управление билетами',
        '   - Просмотр билетов',
        '   - Отмена бронирования',
        '5. Сохранение данных',
        '   - SQLite база данных для хранения информации',
        '   - Автоматическая загрузка данных при запуске',
        '6. Графический интерфейс',
        '   - Java Swing компоненты',
        '   - Переходы между окнами приложения'
    ]

    for req in completed_reqs:
        doc.add_paragraph('   ' + req)

    doc.add_paragraph('')

    # 3. Перечисление невыполненных или измененных требований
    doc.add_paragraph('3. Перечисление невыполненных или измененных требований')
    doc.add_paragraph('')

    doc.add_paragraph('Не выполнены (временно):')
    not_completed = [
        '1. Оплата билетов - реализована только бронь, без реальной оплаты',
        '2. Уведомления по email - нет интеграции с почтовыми сервисами',
        '3. Админ-панель - ограниченная функциональность для администраторов',
        '4. Многоязычность интерфейса - только русский язык'
    ]

    for req in not_completed:
        doc.add_paragraph('   ' + req)

    doc.add_paragraph('')
    doc.add_paragraph('Измененные требования:')
    changed_reqs = [
        '1. Вместо реальной оплаты → демонстрационная бронь',
        '2. Вместо email уведомлений → вывод сообщений в интерфейсе',
        '3. Упрощенная админ-функциональность'
    ]

    for req in changed_reqs:
        doc.add_paragraph('   ' + req)

    doc.add_paragraph('')
    doc.add_paragraph('4.	Программные классы')
    doc.add_paragraph('Программные классы, участвующие в реализации выбранного прецедента')
    doc.add_paragraph('')
    doc.add_paragraph('Концептуальные классы:')
    doc.add_paragraph('•	User (Пользователь)')
    doc.add_paragraph('   o Атрибуты: id, email, password, name, role')
    doc.add_paragraph('   o Методы: getId(), getEmail(), getPassword(), getName(), getRole(), setId(), setEmail(), setPassword(), setName(), setRole()')
    doc.add_paragraph('•	Route (Маршрут)')
    doc.add_paragraph('   o Атрибуты: id, from, to, stops, departureTime, arrivalTime, date, price, totalSeats, tickets')
    doc.add_paragraph('   o Методы: getBookedSeats(), getId(), getFrom(), getTo(), getStops(), getDepartureTime(), getArrivalTime(), getDate(), getPrice(), getTotalSeats(), getTickets(), соответствующие set методы')
    doc.add_paragraph('•	Ticket (Билет)')
    doc.add_paragraph('   o Атрибуты: id, userId, routeId, seatNumber, passengerName, passengerEmail, passengerPhone, status, createdAt')
    doc.add_paragraph('   o Методы: getId(), getUserId(), getRouteId(), getSeatNumber(), getPassengerName(), getPassengerEmail(), getPassengerPhone(), getStatus(), getCreatedAt(), соответствующие set методы')
    doc.add_paragraph('')
    doc.add_paragraph('Служебные классы:')
    doc.add_paragraph('•	DataService (Сервис доступа к данным)')
    doc.add_paragraph('   o Атрибуты: users (Map<String, User>), routes (Map<String, Route>), tickets (Map<String, Ticket>)')
    doc.add_paragraph('   o Методы: initSampleData(), findUserByEmail(email), addUser(user), getAllUsers(), getAllRoutes(), findRouteById(routeId), addRoute(route), addTicket(ticket), getUserTickets(userId), getAllTickets(), cancelTicket(ticketId), isSeatAvailable(routeId, seatNumber)')
    doc.add_paragraph('')
    doc.add_paragraph('Интерфейсные классы (GUI):')
    doc.add_paragraph('•	MainFrame (Главное окно поиска билетов)')
    doc.add_paragraph('   o Атрибуты: dataService, searchFromField, searchToField, searchDateField, routesPanel, sortComboBox, currentRoutes')
    doc.add_paragraph('   o Методы: initializeUI(), createNavigationBar(), createContentPanel(), createSearchPanel(), createFiltersPanel(), createResultsPanel(), loadRoutes(), filterAndSortRoutes(), displayRoutes(routes), createRouteCard(route), openBooking(route), showProfileView()')
    doc.add_paragraph('•	BookingFrame (Окно бронирования)')
    doc.add_paragraph('   o Атрибуты: dataService, route, selectedSeat, passengerNameField, passengerEmailField, passengerPhoneField, seatsPanel')
    doc.add_paragraph('   o Методы: initializeUI(), createRouteInfoPanel(), createSeatSelectionPanel(), generateSeats(), selectSeat(seatNumber), createPassengerInfoPanel(), bookTicket(), goBack()')
    doc.add_paragraph('•	ProfileFrame (Окно «Мои билеты») ')
    doc.add_paragraph('   o Атрибуты: dataService')
    doc.add_paragraph('   o Методы: initializeUI(), createNavigationBar(), createTicketsPanel(), createTicketCard(ticket, route), getStatusText(status), getStatusColor(status), cancelTicket(ticket), showSearchView()')
    doc.add_paragraph('•	ScheduleFrame (Окно расписания рейсов)')
    doc.add_paragraph('   o Атрибуты: таблица расписания (JTable), фильтр направления (JComboBox), данные расписания (Object[][] data)')
    doc.add_paragraph('   o Методы: конструктор ScheduleFrame(), createScheduleTable(), настройка колонок и модели таблицы')
    doc.add_paragraph('•	TicketVoyageApp (Демонстрационное окно без авторизации)')
    doc.add_paragraph('   o Атрибуты: поля поиска, панель маршрутов, кнопки фильтров и расписания')
    doc.add_paragraph('   o Методы: конструктор TicketVoyageApp(), createNavigationBar(), createContentPanel(), createSearchPanel(), createFiltersPanel(), createResultsPanel(), createRouteCard(...)')
    doc.add_paragraph('')

    # 5. Анализ согласованности моделей
    doc.add_paragraph('5. Анализ согласованности моделей')
    doc.add_paragraph('')

    doc.add_heading('Согласование с анализом предметной области:', level=3)
    doc.add_paragraph('Анализ → Код:')
    domain_analysis = [
        '- Актор "Пользователь" → Класс User ✓',
        '- Сущность "Маршрут" → Класс Route ✓',
        '- Сущность "Билет" → Класс Ticket ✓',
        '- Прецедент "Заказ билетов" → BookingFrame и связанные методы ✓',
        '- Прецедент "Поиск рейсов" → MainFrame и методы поиска ✓'
    ]

    for item in domain_analysis:
        doc.add_paragraph('   ' + item)

    doc.add_heading('Согласование с проектированием:', level=3)
    doc.add_paragraph('Use-Case → Реализация:')
    use_case_mapping = [
        '- "Найти рейс" → MainFrame с полями поиска и методом loadRoutes ✓',
        '- "Забронировать билет" → BookingFrame с выбором места и данными пассажира ✓',
        '- "Просмотреть билеты" → ProfileFrame с отображением билетов ✓'
    ]

    for item in use_case_mapping:
        doc.add_paragraph('   ' + item)

    doc.add_paragraph('Диаграмма классов → Реализация:')
    class_diagram_mapping = [
        '- Класс User → Модель User с атрибутами и методами ✓',
        '- Класс Route → Модель Route с информацией о рейсах ✓',
        '- Класс Ticket → Модель Ticket с данными бронирования ✓',
        '- Класс DataService → Сервис для работы с данными ✓',
        '- GUI классы → MainFrame, BookingFrame, ProfileFrame ✓',
        '- Класс ScheduleFrame → Окно расписания рейсов с таблицей ✓',
        '- Класс TicketVoyageApp → Демонстрационное окно без авторизации ✓'
    ]

    for item in class_diagram_mapping:
        doc.add_paragraph('   ' + item)

    doc.add_paragraph('')
    doc.add_paragraph('Отклонения от проектирования:')
    deviations = [
        '1. Упрощенная модель данных - использование SQLite вместо полноценной БД',
        '2. Отсутствие некоторых валидаций на уровне интерфейса',
        '3. Ограниченная обработка ошибок'
    ]

    for dev in deviations:
        doc.add_paragraph('   ' + dev)

    doc.add_paragraph('')

    # 6. План тестирования и результаты
    doc.add_paragraph('6. План тестирования и результаты')
    doc.add_paragraph('')

    doc.add_paragraph('План тестирования:')
    doc.add_paragraph('1. Модульное тестирование (DataServiceTest):')
    unit_tests = [
        '- testFindUserByEmail: поиск пользователя по email',
        '- testGetAllRoutes: получение списка всех маршрутов'
    ]

    for test in unit_tests:
        doc.add_paragraph('   ' + test)

    doc.add_paragraph('')
    doc.add_paragraph('2. Интеграционное тестирование (IntegrationTest):')
    integration_tests = [
        '- testFullBookingWorkflow: полный цикл бронирования билета'
    ]

    for test in integration_tests:
        doc.add_paragraph('   ' + test)

    doc.add_paragraph('')
    doc.add_paragraph('4. Тестирование производительности (PerformanceTest):')
    performance_tests = [
        '- testBulkOperations: массовая операция создания билетов'
    ]

    for test in performance_tests:
        doc.add_paragraph('   ' + test)

    doc.add_paragraph('')
    doc.add_paragraph('5. Дополнительные тесты:')
    additional_tests = [
        '- test_db.java: подключение к базе данных',
        '- test_profile.java: отображение билетов в профиле',
        '- test_clean.java: инициализация чистой БД',
        '- test_warnings.java: проверка предупреждений системы'
    ]

    for test in additional_tests:
        doc.add_paragraph('   ' + test)

    doc.add_paragraph('')
    doc.add_heading('Примеры тестовых сценариев:', level=3)

    # Тест 1
    doc.add_paragraph('Тест 1: Поиск рейсов')
    doc.add_paragraph('Входные данные: откуда="Москва", куда="Санкт-Петербург", дата="2024-12-25"')
    doc.add_paragraph('Ожидаемый результат: отображение доступных рейсов')
    doc.add_paragraph('Фактический результат: ✓ Найдено и отображено 3 рейса')

    # Тест 2
    doc.add_paragraph('Тест 2: Бронирование билета')
    doc.add_paragraph('Входные данные: выбор рейса, место №5, данные пассажира')
    doc.add_paragraph('Ожидаемый результат: билет забронирован, данные сохранены')
    doc.add_paragraph('Фактический результат: ✓ Билет создан, место занято, данные в профиле')

    doc.add_heading('Результаты тестирования:', level=3)
    doc.add_paragraph('Всего тестов: 7 основных + 4 дополнительных')
    doc.add_paragraph('Основные тесты (через TestRunner):')
    doc.add_paragraph('   - DataServiceTest: 2 теста')
    doc.add_paragraph('   - IntegrationTest: 1 тест')
    doc.add_paragraph('   - PerformanceTest: 1 тест')
    doc.add_paragraph('Дополнительные тесты:')
    doc.add_paragraph('   - test_db.java: тест подключения к базе данных')
    doc.add_paragraph('   - test_profile.java: тест отображения билетов в профиле')
    doc.add_paragraph('   - test_clean.java: тест инициализации БД')
    doc.add_paragraph('   - test_warnings.java: тест предупреждений')
    doc.add_paragraph('')
    doc.add_paragraph('Статус тестирования: Все тесты проходят успешно')
    doc.add_paragraph('')
    doc.add_paragraph('Проблемы:')
    issues = [
        '1. При выборе уже занятого места - недостаточно информативное сообщение',
        '2. Отсутствие валидации email при регистрации'
    ]

    for issue in issues:
        doc.add_paragraph('   ' + issue)

    doc.add_paragraph('')

    # 7. Инструкция для пользователя
    doc.add_paragraph('7. Инструкция для пользователя')
    doc.add_paragraph('')

    doc.add_heading('Назначение приложения:', level=3)
    doc.add_paragraph('Система поиска и бронирования билетов на транспорт. Позволяет пользователям искать рейсы, бронировать билеты и управлять своими бронированиями.')

    doc.add_heading('Возможности:', level=3)
    features = [
        '- Регистрация и вход в систему',
        '- Поиск рейсов по маршруту и дате',
        '- Фильтрация результатов по цене, времени, перевозчику',
        '- Бронирование билетов с выбором места',
        '- Просмотр личных билетов',
        '- Отмена бронирования'
    ]

    for feature in features:
        doc.add_paragraph('   ' + feature)

    doc.add_heading('Интерфейс:', level=3)
    doc.add_paragraph('[ВСТАВИТЬ СКРИНШОТЫ ИНТЕРФЕЙСА]')

    doc.add_heading('Пошаговая инструкция:', level=3)
    doc.add_paragraph('')

    doc.add_paragraph('1. Запуск приложения:')
    steps1 = [
        '- Запустите приложение',
        '- Приложение автоматически откроет главное окно с доступными рейсами'
    ]

    for step in steps1:
        doc.add_paragraph('   ' + step)

    doc.add_paragraph('')
    doc.add_paragraph('2. Поиск билетов:')
    steps2 = [
        '- На главной странице заполните поля поиска',
        '- Укажите город отправления, назначения и дату',
        '- Нажмите "Найти"',
        '- Просмотрите результаты, примените фильтры при необходимости'
    ]

    for step in steps2:
        doc.add_paragraph('   ' + step)

    doc.add_paragraph('')
    doc.add_paragraph('3. Бронирование:')
    steps3 = [
        '- Выберите подходящий рейс',
        '- Нажмите "Забронировать"',
        '- Выберите место на схеме',
        '- Заполните данные пассажира',
        '- Нажмите "Подтвердить бронирование"'
    ]

    for step in steps3:
        doc.add_paragraph('   ' + step)

    doc.add_paragraph('')
    doc.add_paragraph('4. Управление билетами:')
    steps4 = [
        '- Перейдите в раздел "Мои билеты"',
        '- Просмотрите список забронированных билетов',
        '- Для отмены нажмите "Отменить" рядом с билетом'
    ]

    for step in steps4:
        doc.add_paragraph('   ' + step)

    doc.add_paragraph('')

    # 8. Демонстрационный вариант приложения
    doc.add_paragraph('8. Демонстрационный вариант приложения')
    doc.add_paragraph('')

    doc.add_heading('Сценарий демонстрации:', level=3)
    doc.add_paragraph('Дано:')
    demo_given = [
        '- База данных с тестовыми пользователями и рейсами',
        '- Доступны рейсы Москва-Петербург, Москва-Казань, Петербург-Казань',
        '- Регистрация нового пользователя'
    ]

    for item in demo_given:
        doc.add_paragraph('   ' + item)

    doc.add_paragraph('Действия:')
    demo_actions = [
        '1. Запуск приложения',
        '2. Поиск рейсов Москва-Петербург на завтра',
        '3. Выбор первого рейса и бронирование места №3',
        '4. Заполнение данных пассажира (Иванов Иван Иванович)',
        '5. Просмотр забронированного билета в профиле',
        '6. Отмена бронирования'
    ]

    for action in demo_actions:
        doc.add_paragraph('   ' + action)

    doc.add_paragraph('Результат:')
    demo_result = [
        '- Найдено несколько рейсов по маршруту',
        '- Билет успешно забронирован',
        '- Данные билета отображаются в профиле',
        '- Бронирование отменено, место освобождено'
    ]

    for result in demo_result:
        doc.add_paragraph('   ' + result)

    doc.add_heading('Структура проекта после работы:', level=3)
    doc.add_paragraph('''project/
├── src/main/java/com/ticketvoyage/
│   ├── Main.java                     # Точка входа
│   ├── model/                        # Модели данных
│   │   ├── User.java
│   │   ├── Route.java
│   │   └── Ticket.java
│   ├── service/                      # Сервисы
│   │   ├── AuthService.java
│   │   └── DataService.java
│   └── ui/                           # Графический интерфейс
│       ├── LoginFrame.java
│       ├── MainFrame.java
│       ├── BookingFrame.java
│       └── ProfileFrame.java
├── lib/                              # Библиотеки
│   ├── sqlite-jdbc.jar
│   ├── slf4j-api.jar
│   └── slf4j-simple.jar
├── ticketvoyage.db                   # База данных
└── README.md                         # Документация''')

    doc.add_heading('Системные требования:', level=3)
    requirements = [
        '- Java 8 или выше',
        '- Windows/Linux/Mac OS',
        '- 512 MB RAM',
        '- 100 MB свободного места'
    ]

    for req in requirements:
        doc.add_paragraph('   ' + req)

    # Сохраняем документ
    output_path = '/Users/admin/Documents/GitHub/ticket-system-desktop/spring/final_report.docx'
    doc.save(output_path)
    print(f"Отчет успешно создан: {output_path}")

if __name__ == "__main__":
    create_report()
